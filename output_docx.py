"""
output_docx.py
DOCX generation using python-docx — professional book-style formatting.
"""

import logging
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────
# Color constants
# ─────────────────────────────────────────────
COLOR_BLACK = RGBColor(0x22, 0x22, 0x22)
COLOR_GRAY = RGBColor(0x66, 0x66, 0x66)


# ─────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────

def _set_page_margins(doc: Document, margin_inches: float = 1.25) -> None:
    """Set uniform page margins on all sections."""
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(margin_inches)
        section.right_margin = Inches(margin_inches)


def _add_page_number(paragraph) -> None:
    """Insert a centered page number field in a paragraph."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def _add_footer_with_page_number(doc: Document) -> None:
    """Add footer with page number to all sections (except first)."""
    for i, section in enumerate(doc.sections):
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        _add_page_number(para)

        # Style the footer text
        for run in para.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)
            run.font.color.rgb = COLOR_GRAY


def _paragraph_spacing(para, space_before_pt: int = 0, space_after_pt: int = 10) -> None:
    """Set paragraph spacing."""
    pf = para.paragraph_format
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)


# ─────────────────────────────────────────────
# DOCX builder
# ─────────────────────────────────────────────

def generate_docx(book, output_path: Path) -> Path:
    """
    Generate a professionally formatted DOCX book.

    Args:
        book: Book dataclass with title, author, chapters list
        output_path: Full path where the DOCX should be saved

    Returns:
        Path to the saved DOCX file
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _set_page_margins(doc, margin_inches=1.25)
    _add_footer_with_page_number(doc)

    # ── Title Page ──────────────────────────────
    # Large vertical spacer
    for _ in range(8):
        spacer = doc.add_paragraph()
        _paragraph_spacing(spacer, 0, 0)

    # Book title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(book.title)
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(32)
    title_run.font.color.rgb = COLOR_BLACK
    _paragraph_spacing(title_para, 0, 24)

    # "by"
    by_para = doc.add_paragraph()
    by_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    by_run = by_para.add_run("by")
    by_run.italic = True
    by_run.font.name = "Times New Roman"
    by_run.font.size = Pt(14)
    by_run.font.color.rgb = COLOR_GRAY
    _paragraph_spacing(by_para, 0, 6)

    # Author name
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_para.add_run(book.author)
    author_run.italic = True
    author_run.font.name = "Times New Roman"
    author_run.font.size = Pt(16)
    author_run.font.color.rgb = COLOR_BLACK
    _paragraph_spacing(author_para, 0, 0)

    # Page break after title page
    doc.add_page_break()

    # ── Chapters ────────────────────────────────
    ORDINALS = [
        "One", "Two", "Three", "Four", "Five",
        "Six", "Seven", "Eight", "Nine", "Ten",
        "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen",
        "Sixteen", "Seventeen", "Eighteen", "Nineteen", "Twenty",
    ]

    for chapter in book.chapters:
        ordinal = (
            ORDINALS[chapter.number - 1]
            if chapter.number <= len(ORDINALS)
            else str(chapter.number)
        )

        # Chapter label
        label_para = doc.add_paragraph()
        label_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        label_run = label_para.add_run(f"CHAPTER {ordinal.upper()}")
        label_run.font.name = "Times New Roman"
        label_run.font.size = Pt(11)
        label_run.font.color.rgb = COLOR_GRAY
        _paragraph_spacing(label_para, 36, 4)

        # Chapter title
        heading_para = doc.add_paragraph()
        heading_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading_run = heading_para.add_run(chapter.title)
        heading_run.bold = True
        heading_run.font.name = "Times New Roman"
        heading_run.font.size = Pt(22)
        heading_run.font.color.rgb = COLOR_BLACK
        _paragraph_spacing(heading_para, 4, 20)

        # Chapter body paragraphs
        paragraphs = [p.strip() for p in chapter.content.split("\n") if p.strip()]
        for i, para_text in enumerate(paragraphs):
            body_para = doc.add_paragraph()
            body_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            body_run = body_para.add_run(para_text)
            body_run.font.name = "Times New Roman"
            body_run.font.size = Pt(12)
            body_run.font.color.rgb = COLOR_BLACK

            # First paragraph of chapter: no indent. Rest: small indent.
            pf = body_para.paragraph_format
            pf.first_line_indent = Inches(0) if i == 0 else Inches(0.3)
            pf.line_spacing = Pt(20)
            _paragraph_spacing(body_para, 0, 8)

        # Page break after each chapter
        doc.add_page_break()

    # Save
    try:
        doc.save(str(output_path))
        logger.info("DOCX saved to: %s", output_path)
    except Exception as exc:
        logger.error("Failed to save DOCX: %s", exc)
        raise RuntimeError(f"DOCX generation failed: {exc}") from exc

    return output_path